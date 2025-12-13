"""
CV Parser Service for extracting content from CV files
"""

import re
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
import io

logger = logging.getLogger(__name__)


class CVParser:
    """CV document parser for PDF and DOCX files"""
    
    def __init__(self):
        """Initialize CV parser"""
        self.supported_formats = ['.pdf', '.doc', '.docx']
        
    async def parse_cv(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse CV content from uploaded file
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Parsed CV data dictionary
        """
        try:
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                return await self._parse_pdf(file_content, filename)
            elif file_extension in ['.doc', '.docx']:
                return await self._parse_docx(file_content, filename)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error parsing CV {filename}: {str(e)}")
            raise
    
    async def _parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse PDF file using pdfplumber for better text extraction"""
        try:
            raw_text = ""
            structured_sections = {}
            
            # Use pdfplumber for better text extraction
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    raw_text += f"\\n--- Page {page_num + 1} ---\\n{page_text}\\n"
            
            # Extract structured information
            structured_sections = self._extract_sections(raw_text)
            contact_info = self._extract_contact_info(raw_text)
            
            return {
                "filename": filename,
                "file_type": "pdf",
                "raw_text": raw_text,
                "structured_sections": structured_sections,
                "contact_info": contact_info,
                "total_pages": len(pdf.pages) if 'pdf' in locals() else 0
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {str(e)}")
            # Fallback to PyPDF2
            return await self._parse_pdf_fallback(file_content, filename)
    
    async def _parse_pdf_fallback(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback PDF parsing using PyPDF2"""
        try:
            raw_text = ""
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                raw_text += f"\\n--- Page {page_num + 1} ---\\n{page_text}\\n"
            
            structured_sections = self._extract_sections(raw_text)
            contact_info = self._extract_contact_info(raw_text)
            
            return {
                "filename": filename,
                "file_type": "pdf",
                "raw_text": raw_text,
                "structured_sections": structured_sections,
                "contact_info": contact_info,
                "total_pages": len(pdf_reader.pages)
            }
            
        except Exception as e:
            logger.error(f"Fallback PDF parsing failed for {filename}: {str(e)}")
            raise
    
    async def _parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            raw_text = ""
            for paragraph in doc.paragraphs:
                raw_text += paragraph.text + "\\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        raw_text += cell.text + " "
                    raw_text += "\\n"
            
            structured_sections = self._extract_sections(raw_text)
            contact_info = self._extract_contact_info(raw_text)
            
            return {
                "filename": filename,
                "file_type": "docx",
                "raw_text": raw_text,
                "structured_sections": structured_sections,
                "contact_info": contact_info,
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {str(e)}")
            raise
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract common CV sections from text"""
        sections = {
            "summary": "",
            "experience": "",
            "education": "", 
            "skills": "",
            "projects": "",
            "certifications": "",
            "languages": ""
        }
        
        # Common section headers patterns
        section_patterns = {
            "summary": r"(?i)(summary|profile|objective|about|overview)\\s*:?\\s*\\n",
            "experience": r"(?i)(experience|work\\s+history|employment|career)\\s*:?\\s*\\n",
            "education": r"(?i)(education|academic|qualifications|studies)\\s*:?\\s*\\n",
            "skills": r"(?i)(skills|technical\\s+skills|competencies|technologies)\\s*:?\\s*\\n",
            "projects": r"(?i)(projects|portfolio|work\\s+samples)\\s*:?\\s*\\n",
            "certifications": r"(?i)(certifications?|licenses?|credentials?)\\s*:?\\s*\\n",
            "languages": r"(?i)(languages?|linguistic\\s+skills?)\\s*:?\\s*\\n"
        }
        
        # Split text into potential sections
        for section_name, pattern in section_patterns.items():
            matches = list(re.finditer(pattern, text, re.MULTILINE))
            if matches:
                start_pos = matches[0].end()
                # Find the end of this section (next section or end of text)
                next_section_pos = len(text)
                for other_pattern in section_patterns.values():
                    if other_pattern != pattern:
                        next_matches = list(re.finditer(other_pattern, text[start_pos:], re.MULTILINE))
                        if next_matches:
                            next_section_pos = min(next_section_pos, start_pos + next_matches[0].start())
                
                sections[section_name] = text[start_pos:next_section_pos].strip()
        
        return sections
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from CV text"""
        contact_info = {
            "name": None,
            "email": None,
            "phone": None,
            "location": None,
            "linkedin": None,
            "github": None
        }
        
        # Email pattern
        email_pattern = r"\\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Z|a-z]{2,}\\b"
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()
        
        # Phone patterns (various formats)
        phone_patterns = [
            r"\\+?\\d{1,4}[\\s.-]?\\(?\\d{1,4}\\)?[\\s.-]?\\d{1,4}[\\s.-]?\\d{1,9}",
            r"\\(?\\d{3}\\)?[\\s.-]?\\d{3}[\\s.-]?\\d{4}",
            r"\\d{3}[\\s.-]?\\d{3}[\\s.-]?\\d{4}"
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                # Clean up the phone number
                phone = re.sub(r"[^\\d+]", "", phone_match.group())
                if len(phone) >= 10:  # Valid phone number
                    contact_info["phone"] = phone_match.group().strip()
                    break
        
        # LinkedIn profile
        linkedin_pattern = r"(?i)linkedin\\.com/in/[\\w-]+"
        linkedin_match = re.search(linkedin_pattern, text)
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()
        
        # GitHub profile
        github_pattern = r"(?i)github\\.com/[\\w-]+"
        github_match = re.search(github_pattern, text)
        if github_match:
            contact_info["github"] = github_match.group()
        
        # Name extraction (usually at the top of CV)
        lines = text.split("\\n")[:10]  # Check first 10 lines
        for line in lines:
            line = line.strip()
            if line and len(line.split()) >= 2 and len(line.split()) <= 4:
                # Potential name (2-4 words, not email or phone)
                if not re.search(r"[@\\d]", line) and not any(keyword in line.lower() for keyword in 
                    ["resume", "cv", "curriculum", "phone", "email", "address"]):
                    contact_info["name"] = line
                    break
        
        return contact_info
    
    def extract_skills_keywords(self, text: str) -> List[str]:
        """Extract potential skill keywords from text"""
        # Common technical skills and technologies
        technical_skills = [
            # Programming languages
            "python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby", "go", "rust",
            "swift", "kotlin", "scala", "r", "matlab", "perl", "shell", "bash",
            
            # Web technologies
            "html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask",
            "spring", "laravel", "rails", "asp.net", "jquery", "bootstrap",
            
            # Databases
            "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "oracle", "sql server",
            "sqlite", "cassandra", "dynamodb",
            
            # Cloud and DevOps
            "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "gitlab", "github",
            "terraform", "ansible", "chef", "puppet",
            
            # Data Science and ML
            "machine learning", "deep learning", "artificial intelligence", "data science",
            "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
            
            # Other tools
            "git", "jira", "confluence", "slack", "teams", "agile", "scrum", "kanban"
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in technical_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return found_skills